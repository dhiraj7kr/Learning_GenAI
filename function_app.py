import logging
import os
import io
import json
import requests
import PyPDF2
import docx
from azure.storage.blob import BlobServiceClient
import azure.functions as func
from langchain.text_splitter import CharacterTextSplitter
from langchain_core.documents import Document
from langchain_openai import AzureOpenAIEmbeddings
from azure.search.documents.indexes.models import (
    SearchableField,
    SearchField,
    SearchFieldDataType,
    SimpleField,
)
from langchain_community.vectorstores.azuresearch import AzureSearch
from openai import AzureOpenAI
from azure.search.documents.models import VectorizedQuery
from azure.search.documents import SearchClient
from azure.core.credentials import AzureKeyCredential

app = func.FunctionApp(http_auth_level=func.AuthLevel.ANONYMOUS)

@app.route(route="http_trigger", methods=["POST"])
def http_trigger(req: func.HttpRequest) -> func.HttpResponse:
    logging.info('Python HTTP trigger function processed a request.')

     # Check if the request method is POST
    if req.method != 'POST':
        return func.HttpResponse(
            json.dumps({"status": "FAILED", "error": "Only POST requests are accepted."}),
            status_code=400,
            mimetype="application/json"
        )
 
    try:
        req_body = req.get_json()
    except ValueError:
        return func.HttpResponse(
            json.dumps({"status": "FAILED", "error": "Invalid JSON body."}),
            status_code=400,
            mimetype="application/json"
        )
 
    doc_link = req_body.get('doc_link')
   
    # Validate document link
    # if not doc_link or not (doc_link.endswith('.pdf') or doc_link.endswith('.docx')):
    #     return func.HttpResponse(
    #         json.dumps({"status": "FAILED", "error": "Invalid document link. Only .pdf and .docx files are accepted."}),
    #         status_code=400,
    #         mimetype="application/json"
    #     )
 
    try:
        # Fetch the document
        response = requests.get(doc_link)
        response.raise_for_status()
        filename=os.path.basename(doc_link.split("?")[0])
    except requests.RequestException as e:
        return func.HttpResponse(
            json.dumps({"status": "FAILED", "error": f"Failed to fetch document: {str(e)}"}),
            status_code=500,
            mimetype="application/json"
        )
 
   
    try:
        # Read the document content
        # if ".pdf" in doc_link:
        if filename.endswith('.pdf'):
            with io.BytesIO(response.content) as file:
                
                reader = PyPDF2.PdfReader(file)
                text = ""
                for page in reader.pages:
                    text += page.extract_text() + "\n"

        else:
            # Read the document content from docx file
            with io.BytesIO(response.content) as file:
                document = docx.Document(file)
                text = ""
                for para in document.paragraphs:
                    text += para.text + "\n"
           
 
    except Exception as e:
        return func.HttpResponse(
            json.dumps({"status": "FAILED", "error": f"Error reading document: {str(e)}"}),
            status_code=500,
            mimetype="application/json"
        )
 
    # Split the content into chunks
    text_splitter = CharacterTextSplitter(
        chunk_size=1000,
        chunk_overlap=200,
        length_function=len,
        is_separator_regex=False
    )
 
    chunks = text_splitter.split_text(text)
 
    embeddings = AzureOpenAIEmbeddings(
        model=str(os.getenv("embedding_model")),
        azure_endpoint=str(os.getenv("AZURE_OPENAI_ENDPOINT")),
        api_key=str(os.getenv("AZURE_OPENAI_API_KEY")),
        api_version="2024-06-01"
    )
   
    vector_store_address: str = str(os.getenv("Azure_search_endpoint"))
    vector_store_password: str = str(os.getenv("Azure_search_key"))
 
    # Set up the vector store
    fields = [
        SearchableField(
            name="id",
            type=SearchFieldDataType.String,
            key=True,
            filterable=True,
            analyzer_name= "keyword"
        ),
        SearchableField(
            name="parent_id",
            type=SearchFieldDataType.String,
            searchable=True,
            filterable=True
        ),
        SearchableField(
            name="content",
            type=SearchFieldDataType.String,
            searchable=True,
            analyzer_name="en.microsoft"
        ),
        SearchField(
            name="content_vector",
            type=SearchFieldDataType.Collection(SearchFieldDataType.Single),
            searchable=True,
            vector_search_dimensions=1536,
            vector_search_profile_name="myHnswProfile",
            # retrievable=True
        ),
        SearchableField(
            name="metadata",
            type=SearchFieldDataType.String,
            searchable=True,
        ),
        # Additional field to store the title
        SearchableField(
            name="title",
            type=SearchFieldDataType.String,
            searchable=True,
        ),
        # Additional field for filtering on document source
        SimpleField(
            name="source",
            type=SearchFieldDataType.String,
            filterable=True,
        ),
    ]
    index_name: str = "3log-index"
 
    vector_store = AzureSearch(
        azure_search_endpoint=vector_store_address,
        azure_search_key=vector_store_password,
        index_name=index_name,
        embedding_function=embeddings.embed_query,
        fields=fields,
    )
 
    # Get embeddings for the chunks and add them to the vector store
    try:
        # vector_query = embeddings.embed_documents(chunks)
    # Prepare the documents to be added to the vector store
        documents_to_add = []
 
        for i, item in enumerate(chunks):
            documents_to_add.append(
                # {
                #     "id": str(i),
                #     "page_content": item,
                #     # "metadata": json.dumps({"source": "3log"})
                # }
                Document(page_content = item,id = i)
            )
        # documents_to_add = [{"id": str(i), "page_content": chunk, "type": 'Document'} for i, chunk in enumerate(chunks)]
        vector_store.add_documents(documents_to_add)
    except Exception as e:
        return func.HttpResponse(
            json.dumps({"status": "FAILED", "error": f"Failed to index document: {str(e)}"}),
            status_code=500,
            mimetype="application/json"
        )
 
 
    return func.HttpResponse(
        json.dumps({"status": "COMPLETED", "error": None}),
        status_code=200,
        mimetype="application/json"
    )
 

@app.route(route="QueryKnowledgeBase", methods=["GET"])
def QueryKnowledgeBase(req: func.HttpRequest) -> func.HttpResponse:
    logging.info('Python HTTP trigger function processed a request.')
    try:
        query = req.params.get('query')
        index_name = req.params.get('index_name')

        if not query and not index_name:
            return func.HttpResponse(
            json.dumps({"response": None, "error": "Two Required Paramas Missing query and index_name" }),
            status_code=500,
            mimetype="application/json"
        )
        elif not query:
            return func.HttpResponse(
            json.dumps({"response": None, "error": "Required Paramas Missing make sure to give query" }),
            status_code=500,
            mimetype="application/json"
        )
        elif not index_name:
            return func.HttpResponse(
            json.dumps({"response": None, "error": "Required Paramas Missing make sure to give index_name" }),
            status_code=500,
            mimetype="application/json"
        )
 
 
        client = AzureOpenAI(
            api_key = os.getenv("AZURE_OPENAI_API_KEY"),  
            api_version = "2024-06-01",
            azure_endpoint =os.getenv("AZURE_OPENAI_ENDPOINT"),
        )
 
        response = client.embeddings.create(
            input = query,
            model= os.getenv("embedding_model"),
            dimensions = 1536
        ).data[0].embedding
       
 
        vector_query = VectorizedQuery(vector=response, k_nearest_neighbors=10, fields="content_vector")
       
        search_client = SearchClient(endpoint=os.getenv("Azure_search_endpoint"),
                        index_name=index_name,
                        credential=AzureKeyCredential(os.getenv("Azure_search_key"))
                        )
       
        results = search_client.search(  
            search_text=None,  
            vector_queries= [vector_query],
            select=["content"],
        )  
       
        formatted_results = ""
        for idx, result in enumerate(results):
            formatted_results += f"{idx + 1}. {result['content']}\n"
       
       
        chat_reponse = client.chat.completions.create(
            model=os.getenv("gpt_model"),
            messages=[
                {
                    "role":"system",
                    "content": f"""{os.getenv("System_prompt")}\n you are assigned with the task to give response if the output or result is available in the {formatted_results}.you are advised to give output for the queries related to the document we have provided, or the Information is being provided.\nIt is strictly advised when user tries to greet, and small talk like ('hi' or 'hello') and all your response should be ###WeDontHaveTheAnswer###"""
                },
                {
                    "role":"user",
                    "content": query
                }
            ],
            temperature=0,
            top_p=0.9,
            max_tokens=300,
        )
       
        chat_completion_result = chat_reponse.choices[0].message.content    
       
        result = True if chat_completion_result!="###WeDontHaveTheAnswer###" else False  # Replace with logic if needed
       
        # Check if there is a valid response
        if result:
            return func.HttpResponse(
                json.dumps({"response": chat_completion_result, "error": None}),
                status_code=200,
                mimetype="application/json"
            )
        else:
            # If no meaningful answer is generated
            return func.HttpResponse(
                json.dumps({"response": None, "error": None}),
                status_code=200,
                mimetype="application/json"
            )
    except Exception as e:
        # Handle unexpected errors
        return func.HttpResponse(
            json.dumps({"response": None, "error": str(e)}),
            status_code=500,
            mimetype="application/json"
        )